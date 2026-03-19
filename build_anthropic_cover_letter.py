import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

def add_para(text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

add_para('Matt Gates')
add_para('Chico, CA | (530) 786-3655 | ridgecellrepair@gmail.com')
add_para('March 18, 2026')
add_para('')
add_para('Anthropic Hiring Team')
add_para('Software Engineer, Business Technology')
add_para('')

add_para('Dear Anthropic Hiring Team,', bold=True)
add_para('')

add_para("I'm applying for the Software Engineer, Business Technology position because it sits at the exact intersection of what I do best and what I care most about: building internal tools that make teams more productive, powered by AI that's safe and reliable.")

add_para("For the past four years, I've been building production internal tools and AI-powered automation systems full-time. My current work includes 10+ MCP (Model Context Protocol) servers in Python and Rust for Claude Code integration, a multi-agent orchestration system that coordinates 7 specialized AI agents, and full-stack dashboards built with Next.js, FastAPI, and PostgreSQL. I build with Claude's API every day \u2014 I understand tool calling, structured output, prompt engineering, and the practical challenges of making LLMs reliable in production environments.")

add_para("What makes me particularly well-suited for this role is that I don't just build AI tools \u2014 I build internal tools that real teams use daily. I've spent years consulting directly with clients to understand their workflows, identifying friction points, and shipping solutions that measurably improve productivity. My OpenLedger project replaced manual bookkeeping with Claude-powered transaction classification. My OpenClaw system automated SEO audits, lead response, and proposal generation for operations teams. In every case, the process was the same: deeply understand the user, prototype rapidly, iterate based on feedback, and ship something reliable.")

add_para("I also bring strong security and infrastructure fundamentals. I manage cloud infrastructure across a Tailscale mesh network with automated failover and monitoring. I've implemented OAuth flows, RBAC, secrets management, and API security patterns in production. I have hands-on experience with AWS, Docker, Kubernetes, CI/CD pipelines, and the operational discipline that comes from running systems that can't go down.")

add_para("The Business Technology team's mission \u2014 enabling Anthropic's teams to work efficiently and securely at scale through thoughtful, easy-to-use tools augmented by Claude \u2014 is exactly the kind of high-leverage work I want to do. I believe AI safety isn't only a research problem; the internal infrastructure that enables your teams to move fast and securely is foundational to shipping safe AI to the world.")

add_para("I'm available immediately, I'm open to working from your San Francisco office as needed, and I'm deeply motivated by Anthropic's mission. I would be honored to contribute to the team building the tools that power Anthropic from the inside.")

add_para('')
add_para('Best regards,')
add_para('Matt Gates')

doc.save(r'C:\Users\Matt\Downloads\matt_gates_cover_letter_anthropic.docx')
print('Cover letter saved: matt_gates_cover_letter_anthropic.docx')
