import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

def center_text(text, size=14, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)

def section_heading(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '4',
        qn('w:space'): '1', qn('w:color'): '000000'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)

def body(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(text)
        run.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)

def role_header(title, dates):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10.5)
    run = p.add_run('    ' + dates)
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(6)

def role_company(company):
    p = doc.add_paragraph()
    run = p.add_run(company)
    run.italic = True
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)

# === HEADER ===
center_text('MATT GATES', size=16)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Chico, CA  |  (530) 786-3655  |  ridgecellrepair@gmail.com')
run.font.size = Pt(9.5)
p.paragraph_format.space_after = Pt(0)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('github.com/suhteevah  |  linkedin.com/in/matt-michels')
run.font.size = Pt(9.5)
p.paragraph_format.space_after = Pt(4)

# === SUMMARY ===
section_heading('PROFESSIONAL SUMMARY')
body('Full-stack engineer with 10 years of experience building production internal tools, automation platforms, and developer infrastructure. Deep expertise in Python, TypeScript, React, and Node.js with production cloud deployment experience on AWS. Currently building AI-powered automation systems using Claude API and MCP (Model Context Protocol), including multi-agent orchestration, tool calling, and structured output pipelines. Passionate about understanding user workflows and building intuitive, reliable tools that increase team productivity.')

# === SKILLS ===
section_heading('TECHNICAL SKILLS')
body('React, Next.js, TypeScript, Node.js, Python, FastAPI, Rust (Axum), Go (familiar)', bold_prefix='Full-Stack: ')
body('Claude API, MCP Servers, Prompt Engineering, AI Agent Architectures, Tool Calling, LangChain, RAG', bold_prefix='AI / LLM: ')
body('AWS (EC2, S3, Lambda), Docker, Kubernetes, CI/CD, Tailscale, Prometheus, Grafana, Linux', bold_prefix='Cloud & DevOps: ')
body('OAuth/OIDC, API Security, RBAC, Secrets Management, SSO Integration', bold_prefix='Security & Auth: ')
body('Test Automation, API Testing, Selenium, Playwright, Regression Suites, CI/CD Integration', bold_prefix='QA / Testing: ')
body('PostgreSQL, SQLite, Redis, REST API Design, GraphQL', bold_prefix='Data & APIs: ')

# === EXPERIENCE ===
section_heading('PROFESSIONAL EXPERIENCE')

role_header('Technical Director / Full-Stack Engineer', '2021 \u2013 Present')
role_company('Ridge Cell Repair LLC  |  Chico, CA (Remote)')
bullet('Built 10+ production MCP servers (Python/Rust) for Claude Code integration, enabling AI-powered fleet management, deployment orchestration, and automated internal workflows')
bullet('Designed and shipped full-stack internal tools: Next.js dashboards, FastAPI backends, PostgreSQL databases, real-time monitoring with Prometheus/Grafana')
bullet('Architected OpenClaw: 7-agent AI orchestration system with central coordinator automating SEO audits, lead response, proposals, and operations \u2014 directly serving internal team needs')
bullet('Built autonomous job-hunting platform: 15-tool MCP server with multi-API search, AI fit scoring, cover letter generation, and SMTP/IMAP email automation')
bullet('Developed OpenLedger: FOSS GAAP-compliant bookkeeping app (FastAPI + PostgreSQL) with Claude AI for automated transaction classification \u2014 internal tool replacing manual processes')
bullet('Managed cloud infrastructure across Tailscale mesh: automated failover, health monitoring, GPU inference deployments with Ollama/Open WebUI')
bullet('Consulted directly with B2B clients to deeply understand workflows, delivered tailored tools and automation that measurably improved operational efficiency')
bullet('Championed API design best practices, security (OAuth, secrets management), and iterative development based on continuous user feedback')

role_header('AI Application Developer / Teaching Assistant', '2022 \u2013 2024')
role_company('Profixerr  |  Remote')
bullet('Led development of AI-powered internal tools: customer service chatbots, technician knowledge bases, and automated support workflows')
bullet('Collaborated directly with end users to gather requirements and iterate on features based on feedback \u2014 focus on usability and reducing friction')
bullet('Built intuitive interfaces that simplified complex technical workflows for non-technical team members')

role_header('Quality Assurance Specialist', 'Mar 2025 \u2013 Present')
role_company('Compac Engineering  |  Paradise, CA')
bullet('QA testing and inspection on industrial sensor systems to OEM specifications \u2014 systematic validation, attention to detail, operational excellence')
bullet('Conducted comprehensive technical audit: identified issues, built improvement roadmap with prioritized action items')

role_header('Store Manager', '2016 \u2013 2021')
role_company('Chico Cell Repair / CPR Chico  |  Chico, CA')
bullet('Managed daily operations and led team of technicians; built internal tracking tools that reduced turnaround time')
bullet('Created customer-facing and internal workflow systems that improved service quality and team productivity')

# === PROJECTS ===
section_heading('KEY PROJECTS')
body('Production Rust MCP server for AI fleet management. Axum, SQLite, Tailscale integration.', bold_prefix='openclaw-admin-mcp \u2014 ')
body('Next.js multi-tenant SaaS dashboard with real-time monitoring, auth, and API integration.', bold_prefix='openclaw-dashboard \u2014 ')
body('Autonomous Python MCP server: 15 tools, multi-API search, AI scoring, email automation.', bold_prefix='job-hunter-mcp \u2014 ')
body('FOSS QuickBooks alternative. FastAPI + PostgreSQL + Claude AI. GAAP-compliant double-entry.', bold_prefix='OpenLedger \u2014 ')
body('Developer security & quality tool suite. 80+ patterns, 6+ languages. Open source.', bold_prefix='TypeDrift / PerfGuard / APIShield \u2014 ')

# === EDUCATION ===
section_heading('EDUCATION')
body('Coursework in Computer Science & Networking', bold_prefix='Butte College \u2014 ')

doc.save(r'C:\Users\Matt\Downloads\matt_gates_resume_anthropic.docx')
print('Resume saved: matt_gates_resume_anthropic.docx')
