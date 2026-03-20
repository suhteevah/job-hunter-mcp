import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Inches

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

def p(text, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)

p('Matt Gates')
p('Chico, CA | (530) 786-3655 | ridgecellrepair@gmail.com')
p('March 19, 2026')
p('')
p('WinCo Foods Hiring Team')
p('Sr. Middleware Developer')
p('')
p('Dear WinCo Foods Hiring Team,', bold=True)
p('')
p("I'm writing to express my strong interest in the Sr. Middleware Developer position. With 10 years of experience building production middleware, integration systems, and full-stack applications, I'm confident I can make an immediate impact on your team.")
p("My background is deeply technical and hands-on. I build production API services and middleware layers daily using Python (FastAPI), Rust (Axum), and Node.js/TypeScript. I've designed and deployed RESTful and GraphQL APIs, message queues, ETL pipelines, and data integration systems that connect multiple backends and services reliably at scale. My current work includes building MCP (Model Context Protocol) servers that orchestrate communication between AI agents, databases, and external APIs \u2014 essentially middleware that coordinates complex multi-system workflows.")
p("I also bring strong DevOps and infrastructure experience. I manage containerized deployments with Docker and Kubernetes, maintain CI/CD pipelines, and monitor production systems with Prometheus and Grafana. I'm experienced with PostgreSQL, SQLite, Redis, and have worked extensively with cloud platforms including AWS. I understand the operational discipline required to keep middleware running reliably \u2014 health checks, automated failover, logging, and alerting are standard practice in everything I build.")
p("What draws me to WinCo is the combination of a technically challenging middleware role with a company that values its employees and operates with integrity. I'm looking for a stable, long-term position where I can contribute meaningful engineering work, and WinCo's employee-owned model and strong reputation make it an ideal fit.")
p("I'm authorized to work in the US, available immediately, and open to relocation if needed. I'd welcome the opportunity to discuss how my middleware and integration experience can benefit your team.")
p('')
p('Best regards,')
p('Matt Gates')

doc.save(r'C:\Users\Matt\Downloads\matt_gates_cover_letter_winco.docx')
print('Cover letter saved: matt_gates_cover_letter_winco.docx')
